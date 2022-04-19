from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ROW_HEIGHT
from docx.shared import Inches,Pt
from docx.oxml.shared import OxmlElement, qn
from docxtpl import DocxTemplate
import datetime
#import dataframe_image as dfi
#import matplotlib
import numpy as np
import pandas as pd
import os


doc = Document()

p1 = doc.add_paragraph('')
run =p1.add_run('STI - Zoliflodacin phase III')
run.italic = False
run.bold = False
font = run.font
font.name = 'Calibri'
font.size = Pt(14)

paragraph_format = p1.paragraph_format
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
paragraph_format.space_before = Pt(48)
paragraph_format.space_after = Pt(18)
p2 = doc.add_paragraph('')
run =p2.add_run('PROGRESS REPORT')
run.italic = False
run.bold = True
font = run.font
font.name = 'Calibri'
font.size = Pt(16)

paragraph_format = p2.paragraph_format
paragraph_format.space_before = Pt(48)
paragraph_format.space_after = Pt(18)
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3 = doc.add_paragraph('By the 5th of every month')
paragraph_format = p3.paragraph_format
paragraph_format.space_before = Pt(48)
paragraph_format.space_after = Pt(4)
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
now = datetime.datetime.now()
p4 = doc.add_paragraph()
run =p4.add_run('(With data received as of '+ now.strftime("%d-%B-%Y")+')')
run.italic = True
run.bold = False
font = run.font
font.name = 'Calibri'
font.size = Pt(10)
paragraph_format = p4.paragraph_format
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_page_break()
p1 = doc.add_paragraph('Table of contents')


paragraph = doc.add_paragraph()
run = paragraph.add_run()
fldChar = OxmlElement('w:fldChar')  # creates a new element
fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
#instrText.text = 'TOC \o "1-3" \h \z \u'  # change 1-3 depending on heading levels you need

fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')
fldChar3 = OxmlElement('w:t')
fldChar3.text = "Right-click to update field."
fldChar2.append(fldChar3)

fldChar4 = OxmlElement('w:fldChar')
fldChar4.set(qn('w:fldCharType'), 'end')

r_element = run._r
r_element.append(fldChar)
r_element.append(instrText)
r_element.append(fldChar2)
r_element.append(fldChar4)
p_element = paragraph._p


doc.add_page_break()
new_section = doc.add_section()
new_section.orientation = WD_ORIENT.LANDSCAPE
doc.save('monthly_report1.docx')
#document.save(PATHS['TEMP_FOLDER']+'monthly_report1.docx')
#os.system('monthly_report1.docx')

doc1 = Document('monthly_report1.docx')
#document.save(PATHS['TEMP_FOLDER']+'monthly_report1.docx')
#doc2 = Document('schedule_template1.docx')
doc2 = Document('ScheduleEvents_template.docx')
sec = doc2.add_section()
sec.orientation = WD_ORIENT.LANDSCAPE
for element in doc2.element.body:
    doc1 .element.body.append(element)
#doc1.add_page_break()

doc1.save('monthly_report2.docx')
#os.system('monthly_report2.docx')

doc = Document()
doc.add_heading('Table 2: a)Subject Enrollment by Site ', level=1)

endata = pd.read_excel("StudySummaryReport.xlsx", skipfooter=10 ,usecols =['Site', 'Enrollment Goal', 'Enrollment Actual'])
# Calculating Percentage
endata['PercentageEnrolled'] = (endata['Enrollment Actual'] /
                  endata['Enrollment Goal']) * 100


endata.PercentageEnrolled = endata.PercentageEnrolled.round(decimals = 0)

endata['PercentageEnrolled'] = endata['PercentageEnrolled'].astype(str) + '%'

table = doc.add_table(rows=endata.shape[0]+1, cols=endata.shape[1], style = "Medium Shading 2 Accent 1")
table_cells = table._cells

# add the header rows.
for j in range(endata.shape[-1]):
    table.cell(0,j).text = endata.columns[j]

# --- add the rest of the data frame ---
for i in range(endata.shape[0]):
    for j, table.cell in enumerate(table.rows[i + 1].cells):
        table.cell.text = str(endata.values[i, j])


doc.add_page_break()

doc.add_heading('b)Weekly Enrollment by Site ', level=1)
endata = pd.read_excel("Medrio_EnrollmentChart_LIVE_STI_Study.xlsx", sheet_name="Data")

table = doc.add_table(rows=endata.shape[0]+1, cols=endata.shape[1], style = "Medium Shading 2 Accent 1")
table_cells = table._cells

# add the header rows.
for j in range(endata.shape[-1]):
    table.cell(0,j).text = endata.columns[j]

# --- add the rest of the data frame ---
for i in range(endata.shape[0]):
    for j, table.cell in enumerate(table.rows[i + 1].cells):
        table.cell.text = str(endata.values[i, j])


doc.save("Medrio_EnrollmentChart_LIVE_STI_Study.docx")
#os.system("Medrio_EnrollmentChart_LIVE_STI_Study.docx")

doc3 = Document('monthly_report2.docx')
doc4 = Document('Medrio_EnrollmentChart_LIVE_STI_Study.docx')
for element in doc4.element.body:
    doc3 .element.body.append(element)
doc3.save('monthly_report3.docx')

doc = Document()
doc.add_page_break()
doc.add_heading('Table 3: Site Forms Summary Report ', level=1)
data = pd.read_excel("Medrio_SiteDataSummaryReport.xlsx", usecols =['Subject', 'Forms Entered', 'Forms Complete', 'Forms Not Expected',
'Forms Not Complete'])

# Calculating forms done
data['FormsExpected'] = data['Forms Complete'] + data['Forms Not Complete']

# Calculating Percentage
data['FormsDone'] = (data['Forms Complete'] /data['FormsExpected']) * 100

data.FormsDone = data.FormsDone.round(decimals = 0)

data['FormsDone'] = data['FormsDone'].astype(str) + '%'

data.rename(columns={'Subject': 'Site Total'}, inplace=True)

data = data.reindex(['Site Total', 'FormsExpected', 'Forms Complete', 'Forms Not Complete', 'Forms Not Expected','FormsDone'], axis=1)



tdata=data[data["Site Total"].str.contains('TOTAL')]

table = doc.add_table(rows=tdata.shape[0]+1, cols=tdata.shape[1], style="Medium Shading 2 Accent 1")
table_cells = table._cells

# add the header rows.
for j in range(tdata.shape[-1]):
    table.cell(0,j).text = tdata.columns[j]

# --- add the rest of the data frame ---
for i in range(tdata.shape[0]):
    for j, table.cell in enumerate(table.rows[i + 1].cells):
        table.cell.text = str(tdata.values[i, j])
doc.save("SiteDataSummaryReport.docx")
#os.system("SiteDataSummaryReport.docx")

doc4 = Document('monthly_report3.docx')
doc5 = Document('SiteDataSummaryReport.docx')
for element in doc5.element.body:
    doc4.element.body.append(element)
doc4.save('monthly_report4.docx')
#os.system('monthly_report4.docx')

doc.add_page_break()
doc = Document()
doc.add_page_break()
doc.add_heading('Table 4: Query Management Status', level=1)
sdata = pd.read_excel("Medrio_SubjectDataSummaryReport.xlsx", usecols =['Site', 'Subject', 'Open Queries', 'Total Queries'])

smdata=sdata.groupby(['Site'])['Open Queries','Total Queries'].agg('sum').reset_index()

table = doc.add_table(rows=smdata.shape[0]+1, cols=smdata.shape[1], style = "Medium Shading 2 Accent 1")
table_cells = table._cells

# add the header rows.
for j in range(smdata.shape[-1]):
    table.cell(0,j).text = smdata.columns[j]

# --- add the rest of the data frame ---
for i in range(smdata.shape[0]):
    for j, table.cell in enumerate(table.rows[i + 1].cells):
        table.cell.text = str(smdata.values[i, j])

p = doc.add_paragraph('')
pf = p.paragraph_format
pf.space_before = Pt(6)
pf.space_after = Pt(0)
run = p.add_run('*Open Queries:Queries pending action.')
run.italic = True
font = run.font
font.name = 'Calibri'
font.size = Pt(10)

doc.save("SubjectDataSummaryReport.docx")
#os.system("SubjectDataSummaryReport.docx")

doc5 = Document('monthly_report4.docx')
doc6 = Document('SubjectDataSummaryReport.docx')
for element in doc6.element.body:
    doc5.element.body.append(element)
doc5.save('STI_Monthly_Report.docx')
os.system('STI_Monthly_Report.docx')




# doc = Document()
# doc.add_page_break()
# #anadata = pd.read_excel("Medrio_EnrollmentChart_LIVE_STI_Study.xlsx", sheet_name="Analysis")
# doc.add_picture('d__medrio_data_temp_Medrio_EnrollmentChart_LIVE_STI_Study.png', height=Inches(8.0), width=Inches(7.0))
# last_paragraph = doc.paragraphs[-1]
# last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
# doc.save("Medrio_EnrollmentChart_LIVE_STI_Study.docx")
# doc7 = Document('STImonthly_report1.docx')
# doc8 = Document('Medrio_EnrollmentChart_LIVE_STI_Study.docx')
# for element in doc8.element.body:
#     doc7.element.body.append(element)
# doc7.save('Medrio_EnrollmentChart.docx')
# #os.system('Medrio_EnrollmentChart.docx')
#
# doc.add_page_break()
# doc = Document()
#
# doc.add_picture('d__medrio_data_temp_Medrio_EnrollmentChart_LIVE_STI_Study.png', height=Inches(8.0), width=Inches(7.0))
# last_paragraph = doc.paragraphs[-1]
# last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
#
# doc.save("Medrio_EnrollmentChart_LIVE_STI_Study1.docx")
# #os.system("Medrio_EnrollmentChart_LIVE_STI_Study1.docx")
#
# doc9 = Document('Medrio_EnrollmentChart.docx')
# doc10 = Document('Medrio_EnrollmentChart_LIVE_STI_Study1.docx')
# for element in doc10.element.body:
#     doc9.element.body.append(element)
# doc9.save('Medrio_EnrollmentChart.docx')
# #os.system('Medrio_EnrollmentChart.docx')
